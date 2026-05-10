"""
Octa KPI — Export Module
Generates EU-formatted Word tables for deliverables, KPIs, milestones
and exports the Gantt chart as SVG (editable in CorelDraw/Canva) and PDF.
"""
import io
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from config import DARK


# ── Table styling helpers ─────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color: str):
    """Set table cell background colour."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    tcPr.append(shd)


def _style_header_row(row, bg_hex: str = "1B2A4A"):
    for cell in row.cells:
        _set_cell_bg(cell, bg_hex)
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold       = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size  = Pt(9)


def _add_header(doc: Document, text: str):
    h = doc.add_heading(text, level=2)
    h.runs[0].font.color.rgb = RGBColor(0x1B, 0x2A, 0x4A)


def _set_col_widths(table, widths_cm: list):
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths_cm):
                cell.width = Cm(widths_cm[i])


# ── Deliverables table (EU format) ───────────────────────────────────────────

def export_deliverables_docx(deliverables: list, proposal: dict) -> bytes:
    """EU-format deliverables table as .docx bytes."""
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(9)

    # Title
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t.add_run(f"List of Deliverables — {proposal.get('acronym','') or proposal.get('proposal_id','')}")
    run.bold = True; run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x1B, 0x2A, 0x4A)
    doc.add_paragraph()

    headers = ["No.", "Deliverable Title", "WP", "Type",
               "Dissemination", "Month", "Status", "Description"]
    widths  = [1.2, 4.5, 1.0, 2.0, 2.2, 1.2, 2.0, 4.5]

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
    _style_header_row(table.rows[0])

    for d in sorted(deliverables, key=lambda x: x.get("deliverable_number","")):
        row = table.add_row().cells
        row[0].text = d.get("deliverable_number","")
        row[1].text = d.get("deliverable_title","")
        row[2].text = str(d.get("wp_id",""))
        row[3].text = d.get("deliverable_type","").replace("_"," ").title()
        row[4].text = d.get("dissemination_level","").title()
        row[5].text = str(d.get("delivery_month",""))
        row[6].text = d.get("status","planned").replace("_"," ").title()
        row[7].text = d.get("deliverable_description","")
        for cell in row:
            for para in cell.paragraphs:
                para.paragraph_format.space_after = Pt(2)

    _set_col_widths(table, widths)
    doc.add_paragraph(f"\nGenerated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ── Milestones table (EU format) ─────────────────────────────────────────────

def export_milestones_docx(milestones: list, proposal: dict) -> bytes:
    doc = Document()
    for section in doc.sections:
        section.left_margin  = Cm(2.5)
        section.right_margin = Cm(2.5)

    style = doc.styles["Normal"]
    style.font.name = "Arial"; style.font.size = Pt(9)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t.add_run(f"List of Milestones — {proposal.get('acronym','') or proposal.get('proposal_id','')}")
    run.bold = True; run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x1B, 0x2A, 0x4A)
    doc.add_paragraph()

    headers = ["No.", "Milestone Title", "WP", "Month",
               "Means of Verification", "Status", "Description"]
    widths  = [1.2, 4.0, 1.0, 1.2, 4.5, 2.0, 4.5]

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
    _style_header_row(table.rows[0])

    for m in sorted(milestones, key=lambda x: x.get("due_month", 0) or 0):
        row = table.add_row().cells
        row[0].text = m.get("milestone_number","")
        row[1].text = m.get("milestone_title","")
        row[2].text = str(m.get("wp_id",""))
        row[3].text = str(m.get("due_month",""))
        row[4].text = m.get("means_of_verification","")
        row[5].text = m.get("status","planned").replace("_"," ").title()
        row[6].text = m.get("milestone_description","")

    _set_col_widths(table, widths)
    doc.add_paragraph(f"\nGenerated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ── KPI table (EU format) ─────────────────────────────────────────────────────

def export_kpis_docx(kpis_short: list, kpis_long: list,
                      proposal: dict) -> bytes:
    doc = Document()
    for section in doc.sections:
        section.left_margin  = Cm(2.5)
        section.right_margin = Cm(2.5)

    style = doc.styles["Normal"]
    style.font.name = "Arial"; style.font.size = Pt(9)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t.add_run(f"Key Performance Indicators — {proposal.get('acronym','') or proposal.get('proposal_id','')}")
    run.bold = True; run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x1B, 0x2A, 0x4A)
    doc.add_paragraph()

    def _add_kpi_table(kpis, section_title):
        _add_header(doc, section_title)
        if not kpis:
            doc.add_paragraph("No KPIs defined.")
            return

        headers = ["No.", "KPI Title", "Level", "Baseline",
                   "Target", "Unit", "Month / Year",
                   "Measurement Method", "Verification Source",
                   "Description"]
        widths = [1.0, 3.5, 2.0, 1.5, 1.5, 1.5, 2.0, 4.0, 4.0, 4.0]

        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        hdr_cells = table.rows[0].cells
        for i, h in enumerate(headers):
            hdr_cells[i].text = h
        _style_header_row(table.rows[0])

        for k in sorted(kpis, key=lambda x: x.get("kpi_number","")):
            row = table.add_row().cells
            row[0].text = k.get("kpi_number","")
            row[1].text = k.get("kpi_title","")
            row[2].text = k.get("kpi_level","").replace("_"," ").title()
            row[3].text = str(k.get("baseline_value",""))
            row[4].text = str(k.get("target_value",""))
            row[5].text = k.get("unit","")
            timing = (str(k.get("target_month","")) if k.get("target_month")
                      else f"Year +{k.get('post_project_year','')}")
            row[6].text = timing
            row[7].text = k.get("measurement_method","")
            row[8].text = k.get("verification_source","")
            row[9].text = k.get("kpi_description","")

        _set_col_widths(table, widths)
        doc.add_paragraph()

    _add_kpi_table(kpis_short, "Short-Term KPIs (During Project)")
    _add_kpi_table(kpis_long,  "Long-Term KPIs (Post-Project Impact)")

    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ── Combined full export ──────────────────────────────────────────────────────

def export_full_docx(structure: dict) -> bytes:
    """Export all tables (deliverables + milestones + KPIs) in one Word document."""
    proposal    = structure.get("proposal") or {}
    deliverables= structure.get("deliverables", [])
    milestones  = structure.get("milestones", [])
    kpis_short  = structure.get("kpis_short", [])
    kpis_long   = structure.get("kpis_long", [])

    doc = Document()
    for section in doc.sections:
        section.left_margin  = Cm(2.5)
        section.right_margin = Cm(2.5)

    style = doc.styles["Normal"]
    style.font.name = "Arial"; style.font.size = Pt(9)

    # Cover
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t.add_run(
        f"Work Plan Tables\n"
        f"{proposal.get('acronym','') or proposal.get('proposal_id','')}"
    )
    run.bold = True; run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x1B, 0x2A, 0x4A)
    doc.add_page_break()

    # ── Deliverables ──────────────────────────────────────────────────────────
    _add_header(doc, "List of Deliverables")
    if deliverables:
        headers = ["No.", "Title", "WP", "Type", "Dissemination", "Month", "Description"]
        table   = doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h
        _style_header_row(table.rows[0])
        for d in sorted(deliverables, key=lambda x: x.get("deliverable_number","")):
            row = table.add_row().cells
            row[0].text = d.get("deliverable_number","")
            row[1].text = d.get("deliverable_title","")
            row[2].text = str(d.get("wp_id",""))
            row[3].text = d.get("deliverable_type","").replace("_"," ").title()
            row[4].text = d.get("dissemination_level","").title()
            row[5].text = str(d.get("delivery_month",""))
            row[6].text = d.get("deliverable_description","")
        _set_col_widths(table, [1.2, 4.0, 1.0, 2.0, 2.2, 1.2, 5.0])
    doc.add_paragraph()

    # ── Milestones ────────────────────────────────────────────────────────────
    doc.add_page_break()
    _add_header(doc, "List of Milestones")
    if milestones:
        headers = ["No.", "Title", "WP", "Month", "Means of Verification", "Description"]
        table   = doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h
        _style_header_row(table.rows[0])
        for m in sorted(milestones, key=lambda x: x.get("due_month",0) or 0):
            row = table.add_row().cells
            row[0].text = m.get("milestone_number","")
            row[1].text = m.get("milestone_title","")
            row[2].text = str(m.get("wp_id",""))
            row[3].text = str(m.get("due_month",""))
            row[4].text = m.get("means_of_verification","")
            row[5].text = m.get("milestone_description","")
    doc.add_paragraph()

    # ── Short-term KPIs ───────────────────────────────────────────────────────
    doc.add_page_break()
    _add_header(doc, "Short-Term KPIs")
    if kpis_short:
        headers = ["No.","Title","Baseline","Target","Unit","Month","Measurement","Description"]
        table   = doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h
        _style_header_row(table.rows[0])
        for k in kpis_short:
            row = table.add_row().cells
            row[0].text = k.get("kpi_number","")
            row[1].text = k.get("kpi_title","")
            row[2].text = str(k.get("baseline_value",""))
            row[3].text = str(k.get("target_value",""))
            row[4].text = k.get("unit","")
            row[5].text = str(k.get("target_month",""))
            row[6].text = k.get("measurement_method","")
            row[7].text = k.get("kpi_description","")

    # ── Long-term KPIs ────────────────────────────────────────────────────────
    doc.add_page_break()
    _add_header(doc, "Long-Term KPIs (Post-Project Impact)")
    if kpis_long:
        headers = ["No.","Title","Baseline","Target","Unit","Year After","Description"]
        table   = doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h
        _style_header_row(table.rows[0])
        for k in kpis_long:
            row = table.add_row().cells
            row[0].text = k.get("kpi_number","")
            row[1].text = k.get("kpi_title","")
            row[2].text = str(k.get("baseline_value",""))
            row[3].text = str(k.get("target_value",""))
            row[4].text = k.get("unit","")
            row[5].text = str(k.get("post_project_year",""))
            row[6].text = k.get("kpi_description","")

    # ── Risk Register ─────────────────────────────────────────────────────────
    risks = structure.get("risks", [])
    if risks:
        doc.add_page_break()
        _add_header(doc, "Risk Register")
        LEVEL_BG = {"critical":"FC8181","high":"F6AD55","medium":"F6CC52","low":"6FCF97"}
        headers  = ["Risk ID","Title","Likelihood","Severity","Risk Level","Related WPs","Mitigation","Status"]
        table    = doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        for i,h in enumerate(headers):
            table.rows[0].cells[i].text = h
        _style_header_row(table.rows[0])
        level_order = {"critical":0,"high":1,"medium":2,"low":3}
        for r in sorted(risks, key=lambda x: level_order.get(x.get("risk_level","medium"),2)):
            level = r.get("risk_level","medium")
            row   = table.add_row().cells
            row[0].text = r.get("risk_number","")
            row[1].text = r.get("risk_title","")
            row[2].text = r.get("likelihood","").title()
            row[3].text = r.get("severity","").title()
            row[4].text = level.title()
            row[5].text = r.get("_wp_labels","")
            row[6].text = r.get("mitigation_strategy","")
            row[7].text = r.get("status","").title()
            _set_cell_bg(row[4], LEVEL_BG.get(level,"FFFFFF"))
            for run in row[4].paragraphs[0].runs:
                run.bold = True

    doc.add_paragraph(f"\n\nGenerated by Octa Platform · {datetime.now().strftime('%Y-%m-%d')}")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ── Risk Table (EU format) ────────────────────────────────────────────────────

def export_risks_docx(risks: list, proposal: dict,
                       wp_map: dict = None) -> bytes:
    """
    EU-format risk register table as .docx bytes.
    wp_map: {wp_id: wp_number} for displaying related WPs.
    """
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(9)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t.add_run(
        f"Risk Register — {proposal.get('acronym','') or proposal.get('proposal_id','')}"
    )
    run.bold = True; run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x1B, 0x2A, 0x4A)
    doc.add_paragraph()

    # Risk level colour map (Word highlight approximation using cell bg)
    LEVEL_BG = {
        "critical": "FC8181",
        "high":     "F6AD55",
        "medium":   "F6CC52",
        "low":      "6FCF97",
    }

    headers = ["Risk ID", "Risk Title", "Description",
               "Likelihood", "Severity", "Risk Level",
               "Related WPs", "Mitigation Strategy", "Status"]
    widths  = [1.2, 3.0, 4.0, 1.8, 1.8, 1.8, 2.0, 5.0, 1.8]

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
    _style_header_row(table.rows[0])

    level_order = {"critical":0,"high":1,"medium":2,"low":3}
    risks_sorted= sorted(risks, key=lambda x: level_order.get(x.get("risk_level","medium"),2))

    for r in risks_sorted:
        level   = r.get("risk_level","medium")
        bg_hex  = LEVEL_BG.get(level, "FFFFFF")

        row = table.add_row().cells
        row[0].text = r.get("risk_number","")
        row[1].text = r.get("risk_title","")
        row[2].text = r.get("risk_description","")
        row[3].text = r.get("likelihood","").title()
        row[4].text = r.get("severity","").title()
        row[5].text = level.title()
        row[6].text = r.get("_wp_labels","")   # pre-resolved below
        row[7].text = r.get("mitigation_strategy","")
        row[8].text = r.get("status","").title()

        # Colour the Risk Level cell
        _set_cell_bg(row[5], bg_hex)
        for run in row[5].paragraphs[0].runs:
            run.bold = True

    _set_col_widths(table, widths)
    doc.add_paragraph(f"\nGenerated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ── Budget Excel Export (EU format) ──────────────────────────────────────────

def export_budget_excel(proposal: dict, work_packages: list,
                         budget_entries: list, all_partners: list) -> bytes:
    """
    EU-format budget Excel file.
    One sheet per partner:
      - Header with proposal + partner info
      - For each WP: rows per cost category, then WP subtotal
      - Partner summary table (totals per category across all WPs)
      - Grand total row
    """
    from openpyxl import Workbook
    from openpyxl.styles import (Font, PatternFill, Alignment, Border, Side)
    from openpyxl.utils import get_column_letter

    # ── Colour palette ────────────────────────────────────────────────────────
    C_HEADER_FILL  = "1B2A4A"   # dark navy  → proposal/partner title rows
    C_WP_FILL      = "2E4A7A"   # mid blue   → WP header rows
    C_CAT_FILL     = "EBF1F8"   # very light blue → cost category rows
    C_TOTAL_FILL   = "00BCD4"   # teal       → WP total rows
    C_SUMMARY_FILL = "0097A7"   # darker teal → summary section header
    C_GRAND_FILL   = "1B2A4A"   # navy       → grand total row
    C_ALT_FILL     = "F7FAFC"   # near-white → alternating rows
    WHITE          = "FFFFFF"

    # ── Category labels & ordering ────────────────────────────────────────────
    CAT_ORDER = [
        ("personnel",           "1. Personnel Costs"),
        ("travel",              "2. Travel & Subsistence"),
        ("equipment",           "3. Equipment & Infrastructure"),
        ("other_goods_services","4. Other Goods & Services"),
        ("subcontracting",      "5. Subcontracting"),
        ("third_parties",       "6. Fund for Third Parties"),
        ("overhead",            "7. Overhead / Indirect Costs"),
    ]
    CAT_LABEL = dict(CAT_ORDER)

    # ── Border style ──────────────────────────────────────────────────────────
    thin = Side(border_style="thin", color="BBCCDD")
    thick= Side(border_style="medium",color="1B2A4A")
    def _border(top=None,bottom=None,left=None,right=None):
        return Border(top=top,bottom=bottom,left=left,right=right)

    # ── Helper: apply fill ────────────────────────────────────────────────────
    def _fill(hex_color):
        return PatternFill("solid", fgColor=hex_color)

    def _font(bold=False, color="1A1A2E", size=10, italic=False):
        return Font(bold=bold, color=color, size=size, italic=italic,
                    name="Calibri")

    def _align(h="left", v="center", wrap=False):
        return Alignment(horizontal=h, vertical=v, wrap_text=wrap)

    def _set_row(ws, row_num, values, fill_hex=None, bold=False,
                 font_color="1A1A2E", font_size=10, align_h="left",
                 border_side=None, number_fmt=None, italic=False):
        """Write a row and apply formatting."""
        for col_i, val in enumerate(values, 1):
            cell = ws.cell(row=row_num, column=col_i, value=val)
            if fill_hex:
                cell.fill = _fill(fill_hex)
            cell.font = _font(bold=bold, color=font_color,
                              size=font_size, italic=italic)
            cell.alignment = _align(
                h="right" if col_i >= 3 else align_h, v="center"
            )
            if border_side:
                cell.border = border_side
            if number_fmt and col_i >= 3 and isinstance(val, (int,float)):
                cell.number_format = number_fmt

    def _merge_and_write(ws, row_num, start_col, end_col, value,
                          fill_hex=None, bold=False, font_color="1A1A2E",
                          font_size=10, align_h="left"):
        ws.merge_cells(
            start_row=row_num, start_column=start_col,
            end_row=row_num,   end_column=end_col
        )
        cell = ws.cell(row=row_num, column=start_col, value=value)
        if fill_hex:
            cell.fill = _fill(fill_hex)
        cell.font  = _font(bold=bold, color=font_color, size=font_size)
        cell.alignment = _align(h=align_h, v="center")

    # ── Build lookup structures ────────────────────────────────────────────────
    partners_map = {p["id"]: p for p in all_partners}
    wps_map      = {w["wp_id"]: w for w in work_packages}
    wps_sorted   = sorted(work_packages, key=lambda w: w.get("wp_number",""))

    # Group budget entries: {partner_id: {wp_id: {category: {planned, spent}}}}
    budget: dict = {}
    for e in budget_entries:
        pid  = e.get("partner_id")
        wid  = e.get("wp_id")
        cat  = e.get("cost_category","")
        plan = float(e.get("planned_amount",0) or 0)
        spent= float(e.get("actual_spent_amount",0) or 0)
        if pid not in budget:
            budget[pid] = {}
        if wid not in budget[pid]:
            budget[pid][wid] = {}
        budget[pid][wid][cat] = {"planned": plan, "spent": spent}

    acronym = proposal.get("acronym","") or proposal.get("proposal_id","")

    # ── Workbook ──────────────────────────────────────────────────────────────
    wb = Workbook()
    wb.remove(wb.active)    # remove default empty sheet

    # ── One sheet per partner ─────────────────────────────────────────────────
    for pid, wp_data in budget.items():
        partner = partners_map.get(pid, {})
        pname   = partner.get("full_name","") or partner.get("short_name","") or f"Partner {pid}"
        short   = (partner.get("short_name","") or pname)[:28]

        # Sheet name: max 31 chars, no special chars
        safe_name = "".join(c for c in short if c not in r'\/*?[]:'  )[:31]
        ws = wb.create_sheet(title=safe_name or f"P{pid}")

        # Column widths
        ws.column_dimensions["A"].width = 32
        ws.column_dimensions["B"].width = 28
        ws.column_dimensions["C"].width = 18
        ws.column_dimensions["D"].width = 16
        ws.column_dimensions["E"].width = 16

        row = 1

        # ── Sheet title ───────────────────────────────────────────────────────
        _merge_and_write(ws, row, 1, 5,
                         f"EU Budget Table — {acronym}",
                         fill_hex=C_HEADER_FILL, bold=True,
                         font_color=WHITE, font_size=12)
        ws.row_dimensions[row].height = 22
        row += 1

        _merge_and_write(ws, row, 1, 5,
                         f"Partner: {pname}",
                         fill_hex=C_HEADER_FILL, bold=True,
                         font_color=WHITE, font_size=11)
        ws.row_dimensions[row].height = 20
        row += 1
        row += 1   # blank row

        # ── Column headers ────────────────────────────────────────────────────
        col_headers = ["Cost Category", "Description",
                       "Planned (EUR)", "Spent (EUR)", "Remaining (EUR)"]
        _set_row(ws, row, col_headers,
                 fill_hex=C_SUMMARY_FILL, bold=True, font_color=WHITE,
                 font_size=10,
                 border_side=_border(bottom=thick))
        ws.row_dimensions[row].height = 18
        row += 1

        # ── Per-WP breakdown ──────────────────────────────────────────────────
        partner_planned = 0.0
        partner_spent   = 0.0
        summary_by_cat: dict = {c: {"planned":0.0,"spent":0.0} for c,_ in CAT_ORDER}

        alt = False
        for wp in wps_sorted:
            wid = wp["wp_id"]
            if wid not in wp_data:
                continue   # this partner has no budget in this WP

            wp_num   = wp.get("wp_number","")
            wp_title = wp.get("wp_title","")

            # WP header row
            _merge_and_write(ws, row, 1, 5,
                             f"{wp_num}: {wp_title}",
                             fill_hex=C_WP_FILL, bold=True,
                             font_color=WHITE, font_size=10)
            ws.row_dimensions[row].height = 17
            row += 1

            wp_planned = 0.0
            wp_spent   = 0.0

            for cat_key, cat_label in CAT_ORDER:
                entry   = wp_data[wid].get(cat_key,{})
                planned = entry.get("planned",0.0)
                spent   = entry.get("spent",0.0)
                if planned == 0 and spent == 0:
                    continue    # skip empty categories

                remaining = planned - spent
                row_fill  = C_ALT_FILL if alt else None
                alt       = not alt

                _set_row(ws, row,
                         [cat_label, "", planned, spent, remaining],
                         fill_hex=row_fill, bold=False,
                         number_fmt='#,##0.00 "€"',
                         border_side=_border(bottom=_border(bottom=thin).bottom))
                ws.row_dimensions[row].height = 15
                row += 1

                wp_planned += planned
                wp_spent   += spent
                summary_by_cat[cat_key]["planned"] += planned
                summary_by_cat[cat_key]["spent"]   += spent

            # WP subtotal row
            wp_remaining = wp_planned - wp_spent
            total_cells = [f"{wp_num} — TOTAL", "", wp_planned, wp_spent, wp_remaining]
            _set_row(ws, row, total_cells,
                     fill_hex=C_TOTAL_FILL, bold=True, font_color=WHITE,
                     number_fmt='#,##0.00 "€"',
                     border_side=_border(top=thick, bottom=thick))
            ws.row_dimensions[row].height = 17
            row += 1
            row += 1   # spacer

            partner_planned += wp_planned
            partner_spent   += wp_spent

        # ── Summary section ───────────────────────────────────────────────────
        row += 1
        _merge_and_write(ws, row, 1, 5,
                         f"BUDGET SUMMARY — {short}",
                         fill_hex=C_SUMMARY_FILL, bold=True,
                         font_color=WHITE, font_size=11)
        ws.row_dimensions[row].height = 20
        row += 1

        # Summary column headers
        _set_row(ws, row,
                 ["Cost Category", "", "Total Planned", "Total Spent", "Remaining"],
                 fill_hex=C_WP_FILL, bold=True, font_color=WHITE,
                 border_side=_border(bottom=thick))
        ws.row_dimensions[row].height = 17
        row += 1

        for cat_key, cat_label in CAT_ORDER:
            totals = summary_by_cat.get(cat_key,{})
            plan_t = totals.get("planned",0.0)
            spent_t= totals.get("spent",0.0)
            if plan_t == 0 and spent_t == 0:
                continue
            rem_t  = plan_t - spent_t
            _set_row(ws, row,
                     [cat_label, "", plan_t, spent_t, rem_t],
                     fill_hex=C_CAT_FILL, number_fmt='#,##0.00 "€"')
            ws.row_dimensions[row].height = 15
            row += 1

        # Grand total
        grand_rem = partner_planned - partner_spent
        _set_row(ws, row,
                 ["GRAND TOTAL", "", partner_planned, partner_spent, grand_rem],
                 fill_hex=C_GRAND_FILL, bold=True, font_color=WHITE,
                 font_size=11, number_fmt='#,##0.00 "€"',
                 border_side=_border(top=thick, bottom=thick))
        ws.row_dimensions[row].height = 20
        row += 1

        # ── Freeze top rows ───────────────────────────────────────────────────
        ws.freeze_panes = "A5"

        # ── Print settings ────────────────────────────────────────────────────
        ws.page_setup.orientation = "landscape"
        ws.page_setup.paperSize   = ws.PAPERSIZE_A4
        ws.page_setup.fitToPage   = True
        ws.page_setup.fitToWidth  = 1
        ws.page_setup.fitToHeight = 0

    # ── Summary sheet (all partners) ─────────────────────────────────────────
    ws_all = wb.create_sheet(title="📊 All Partners", index=0)
    ws_all.column_dimensions["A"].width = 35
    ws_all.column_dimensions["B"].width = 18
    ws_all.column_dimensions["C"].width = 16
    ws_all.column_dimensions["D"].width = 16

    row_a = 1
    _merge_and_write(ws_all, row_a, 1, 4,
                     f"Budget Overview — {acronym}",
                     fill_hex=C_HEADER_FILL, bold=True,
                     font_color=WHITE, font_size=13)
    ws_all.row_dimensions[row_a].height = 24
    row_a += 1
    row_a += 1

    _set_row(ws_all, row_a,
             ["Partner", "Planned (EUR)", "Spent (EUR)", "Remaining (EUR)"],
             fill_hex=C_SUMMARY_FILL, bold=True, font_color=WHITE)
    ws_all.row_dimensions[row_a].height = 17
    row_a += 1

    grand_all_plan = 0.0; grand_all_spent = 0.0
    alt = False
    for pid, wp_data in budget.items():
        p      = partners_map.get(pid,{})
        pname  = p.get("full_name","") or f"Partner {pid}"
        plan_p = sum(
            e.get("planned",0.0)
            for wid in wp_data.values()
            for e in wid.values()
        )
        spent_p= sum(
            e.get("spent",0.0)
            for wid in wp_data.values()
            for e in wid.values()
        )
        rem_p  = plan_p - spent_p
        row_fill = C_CAT_FILL if alt else None
        alt = not alt
        _set_row(ws_all, row_a,
                 [pname, plan_p, spent_p, rem_p],
                 fill_hex=row_fill, number_fmt='#,##0.00 "€"')
        ws_all.row_dimensions[row_a].height = 15
        row_a += 1
        grand_all_plan  += plan_p
        grand_all_spent += spent_p

    _set_row(ws_all, row_a,
             ["TOTAL", grand_all_plan, grand_all_spent, grand_all_plan-grand_all_spent],
             fill_hex=C_GRAND_FILL, bold=True, font_color=WHITE,
             number_fmt='#,##0.00 "€"')
    ws_all.row_dimensions[row_a].height = 20

    ws_all.freeze_panes = "A4"

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf.read()
